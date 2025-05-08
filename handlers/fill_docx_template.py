from docx import Document


def fill_docx_template(data: dict, template_path: str, output_path: str):
    doc = Document(template_path)

    def replace_placeholders(text):
        for key, value in data.items():
            text = text.replace(f"{{{{{key}}}}}", value)
        return text

    for paragraph in doc.paragraphs:
        paragraph.text = replace_placeholders(paragraph.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.text = replace_placeholders(cell.text)

    doc.save(output_path)
